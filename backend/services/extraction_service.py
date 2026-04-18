"""
Document Text Extraction Service
Supports: PDF, DOC, DOCX, JPG, JPEG, PNG, BMP, TIFF
"""
import os

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".jpg", ".jpeg", ".png", ".bmp", ".tiff"}


def is_allowed(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS


def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber with PyPDF2 as fallback."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                content = page.extract_text()
                if content:
                    text += f"\n\n[--- PAGE {i+1} ---]\n{content}\n"
    except Exception:
        # Fallback: PyPDF2
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                content = page.extract_text()
                if content:
                    text += f"\n\n[--- PAGE {i+1} ---]\n{content}\n"
    return text.strip()


def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_from_image(file_path: str) -> str:
    """Extract text from image using Tesseract OCR."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        # Improve OCR accuracy: convert to grayscale
        img = img.convert("L")
        text = pytesseract.image_to_string(img, lang="eng")
        return text.strip()
    except pytesseract.TesseractNotFoundError:
        raise ValueError(
            "Tesseract OCR is not installed. "
            "Please install it from: https://github.com/tesseract-ocr/tesseract "
            "and add it to your system PATH."
        )


def extract_text(file_path: str, filename: str) -> str:
    """
    Route extraction based on file extension.
    Returns extracted plain text.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text = extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = extract_from_docx(file_path)
    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff"):
        text = extract_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOC, DOCX, JPG, PNG")

    if not text.strip():
        raise ValueError("Could not extract any readable text from the document.")

    return text
